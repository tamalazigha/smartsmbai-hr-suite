"""
report_agent.py — Agent 6: Shortlist Report (SmartSMBAI v2)
Full port of all IncepifyAI upgrades adapted for SmartSMBAI.
Key differences: region-first, risks_or_concerns/recommended_follow_up_questions columns,
certification gate, commission model understanding check, SmartSMBAI branding.
"""
import os, json
from datetime import datetime, timezone
from io import BytesIO
from dotenv import load_dotenv
import anthropic
from database import _sb, _log, update_candidate_status, save_human_review

load_dotenv()
client=anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
MODEL="claude-sonnet-4-6"
SYSTEM_PROMPT=open("system_prompt_report.txt").read()

NAVY=(26,43,94); BLUE=(37,99,235); TEAL=(5,150,105)
GREEN=(22,101,52); AMBER=(217,119,6); RED=(220,38,38)
GRAY=(107,114,128); LGRAY=(243,244,246); DGRAY=(17,24,39)

MAX_MEMO_CANDIDATES=10

# ── Database helpers ────────────────────────────────────────────────
def get_shortlisted_candidates(region=None, candidate_ids=None) -> list:
    """Flat 4-step query — avoids nested select truncation for 3+ candidates."""
    if not _sb: return []
    try:
        q=_sb.table("candidates").select("*").in_("status",["shortlisted","offered","certified"])
        if region: q=q.eq("region",region)
        if candidate_ids: q=q.in_("id",candidate_ids)
        candidates=q.execute().data or []
        if not candidates: return []
        cids=[c["id"] for c in candidates]
        sess_res=(_sb.table("interview_sessions").select("*")
                  .in_("candidate_id",cids).order("created_at",desc=True).execute())
        session_map={}
        for s in (sess_res.data or []):
            cid=s.get("candidate_id","")
            if cid and cid not in session_map: session_map[cid]=s
        sids=[s["id"] for s in session_map.values() if s.get("id")]
        score_map={}
        if sids:
            sc_res=(_sb.table("interview_scores").select("*")
                    .in_("session_id",sids).order("scored_at",desc=True).execute())
            for sc in (sc_res.data or []):
                sid=sc.get("session_id","")
                if sid and sid not in score_map: score_map[sid]=sc
        rev_res=(_sb.table("human_reviews").select("*")
                 .in_("candidate_id",cids).order("reviewed_at",desc=True).execute())
        review_map={}
        for r in (rev_res.data or []):
            cid=r.get("candidate_id","")
            if cid: review_map.setdefault(cid,[]).append(r)
        enriched=[]
        for c in candidates:
            cid=c["id"]; session=session_map.get(cid,{}); sid=session.get("id","")
            enriched.append({**c,"_session":session,"_score":score_map.get(sid,{}),"_reviews":review_map.get(cid,[])})
        return enriched
    except Exception as e: print(f"[Report] get_shortlisted: {e}"); return []

def get_offered_candidates(region=None) -> list:
    if not _sb: return []
    try:
        q=_sb.table("candidates").select("*, human_reviews(*)").in_("status",["offered","certified"])
        if region: q=q.eq("region",region)
        return q.execute().data or []
    except Exception as e: print(f"[Report] get_offered: {e}"); return []

def get_candidate_cv(candidate_id: str) -> tuple:
    if not _sb: return "",""
    try:
        res=(_sb.table("applications").select("cv_text,cover_letter_text")
             .eq("candidate_id",candidate_id).order("received_at",desc=True).limit(1).execute())
        if res.data:
            row=res.data[0]; return row.get("cv_text","") or "",row.get("cover_letter_text","") or ""
        return "",""
    except Exception as e: print(f"[Report] get_cv: {e}"); return "",""

def mark_as_offered(candidate_id: str, reviewer: str, notes: str="") -> bool:
    try:
        update_candidate_status(candidate_id,"offered",actor=reviewer)
        save_human_review(candidate_id,reviewer,"shortlist","advance",notes or f"Offer extended by {reviewer}")
        _log("offer_extended",actor=reviewer,object_type="candidate",object_id=candidate_id,
             summary=f"Offer extended by {reviewer}")
        return True
    except Exception as e: print(f"[Report] mark_as_offered: {e}"); return False

def get_report_history(candidate_ids: list) -> dict:
    if not _sb or not candidate_ids: return {}
    try:
        res=(_sb.table("hr_audit_log").select("object_id,timestamp_utc,actor")
             .eq("action","candidate_in_report").in_("object_id",candidate_ids)
             .order("timestamp_utc",desc=True).execute())
        history={}
        for row in (res.data or []):
            cid=str(row.get("object_id",""))
            if cid: history.setdefault(cid,[]).append({"at":(row.get("timestamp_utc") or "")[:16].replace("T"," "),"by":row.get("actor","HR")})
        return history
    except Exception as e: print(f"[Report] get_report_history: {e}"); return {}

def log_candidates_in_report(candidate_ids: list, prepared_by: str, region: str="") -> None:
    for cid in candidate_ids:
        _log("candidate_in_report",actor=prepared_by,object_type="candidate",object_id=cid,
             summary=f"Included in shortlist report ({region}) by {prepared_by}")

# ── Memo generation ─────────────────────────────────────────────────
def _build_candidate_summary(c: dict) -> dict:
    score_data=c.get("_score") or {}
    reviews=c.get("_reviews") or []
    if not score_data:
        sessions=c.get("interview_sessions",[])
        if sessions:
            sessions_sorted=sorted(sessions,key=lambda s:s.get("created_at",""),reverse=True)
            sc_list=sessions_sorted[0].get("interview_scores",[])
            if sc_list: score_data=sc_list[0]
    if not reviews: reviews=c.get("human_reviews",[])
    latest=reviews[0] if reviews else {}
    dim_keys=["sales_track_record","local_network_proof","tech_ai_prompting_readiness",
              "ai_troubleshooting_integration_privacy","ai_adoption_metrics_thinking",
              "communication_clarity","region_fit"]
    scores={k:score_data.get(k,0) for k in dim_keys}
    return {"id":c.get("id"),"name":c.get("name","Unknown"),"email":c.get("email",""),
            "region":c.get("region","—"),"total_score":score_data.get("total_score",0),
            "scores":scores,"recommendation":score_data.get("recommendation") or "—",
            "summary":score_data.get("summary") or "",
            "evidence":score_data.get("evidence_highlights") or [],
            "risks":score_data.get("risks_or_concerns") or [],
            "follow_up":score_data.get("recommended_follow_up_questions") or [],
            "human_decision":latest.get("decision") or "—",
            "human_notes":latest.get("notes") or "—",
            "certified":c.get("status")=="certified",
            "_has_score":bool(score_data)}

def _format_candidate_for_prompt(s: dict, idx: int) -> str:
    scores=s.get("scores",{})
    dim_abbr={"sales_track_record":"Sales","local_network_proof":"Network",
               "tech_ai_prompting_readiness":"Tech/AI","ai_troubleshooting_integration_privacy":"Troubleshoot",
               "ai_adoption_metrics_thinking":"Adoption","communication_clarity":"Comms","region_fit":"Region"}
    dim_str=" | ".join(f"{dim_abbr.get(k,k[:6])}:{v}" for k,v in scores.items()) if scores else "(no scores)"
    evidence=s.get("evidence",[]); risks=s.get("risks",[]); follow_up=s.get("follow_up",[])
    missing_note="  WARNING: No score data for this candidate.\n" if not s.get("_has_score",True) else ""
    lines=[f"CANDIDATE {idx}: {s['name']} ({s['email']}) — Region: {s['region']}",
           missing_note if missing_note else "",
           f"  Score: {s['total_score']}/35 | AI Rec: {s['recommendation']} | Certified: {s.get('certified',False)}",
           f"  Dimensions: {dim_str}",
           f"  AI Summary: {(s.get('summary') or '(none)')[:500]}",
           f"  Evidence: {' | '.join(f'[{i+1}] {e}' for i,e in enumerate(evidence)) or '(none)'}",
           f"  Risks: {' | '.join(f'[{i+1}] {r}' for i,r in enumerate(risks)) or '(none)'}",
           f"  Follow-up: {' | '.join(f'[{i+1}] {q}' for i,q in enumerate(follow_up)) or '(none)'}",
           f"  HR Decision: {s['human_decision']} | HR Notes: {s['human_notes']}"]
    return "\n".join(l for l in lines if l)

def _validate_summaries(summaries: list) -> list:
    warnings=[]
    for s in summaries:
        if not s.get("_has_score",True):
            warnings.append(f"{s.get('name','?')} ({s.get('region','?')}): no score data found.")
        elif s.get("total_score",0)==0:
            warnings.append(f"{s.get('name','?')}: total score is 0 — may be incomplete.")
    return warnings

def _single_candidate_memo(candidates: list, region: str, prepared_by: str) -> dict:
    s=_build_candidate_summary(candidates[0])
    log_candidates_in_report([s["id"]] if s.get("id") else [],prepared_by,region)
    _log("report_generated",actor=prepared_by,object_type="region",summary=f"Single-candidate — {region} — {s['name']}")
    return {"region":region,"prepared_by":prepared_by,"generated_at":datetime.now(timezone.utc).isoformat(),
            "candidate_count":1,"single_candidate":True,
            "executive_summary":(f"{s['name']} is the only shortlisted candidate for {region}. "
                                 f"Score: {s['total_score']}/35. AI rec: {s.get('recommendation','—')}. "
                                 f"Certification: {'complete' if s.get('certified') else 'pending'}."),
            "ranked_candidates":[{"rank":1,"name":s["name"],"email":s["email"],"region":s["region"],
                                  "score":s["total_score"],"key_strength":s["evidence"][0] if s["evidence"] else "See scorecard",
                                  "key_concern":s["risks"][0] if s["risks"] else "See scorecard",
                                  "follow_up_questions":s["follow_up"],"commission_model_understanding":"—",
                                  "summary":s["summary"],"evidence":s["evidence"],"risks":s["risks"],
                                  "certified":s.get("certified",False)}],
            "recommended_first_hire":f"{s['name']} — sole shortlisted candidate.",
            "live_interview_guidance":s["follow_up"],"certification_status":[],
            "process_notes":"Only one candidate shortlisted."}

def _multi_candidate_memo(candidates: list, region: str, prepared_by: str) -> dict:
    summaries=sorted([_build_candidate_summary(c) for c in candidates],
                     key=lambda x:x.get("total_score",0),reverse=True)
    excluded=[]
    if len(summaries)>MAX_MEMO_CANDIDATES:
        excluded=summaries[MAX_MEMO_CANDIDATES:]; summaries=summaries[:MAX_MEMO_CANDIDATES]
    warnings=_validate_summaries(summaries)
    for w in warnings: print(f"[Report] Warning: {w}")
    candidates_text="\n\n".join(_format_candidate_for_prompt(s,i+1) for i,s in enumerate(summaries))
    warnings_text=("\n\nDATA WARNINGS:\n"+"".join(f"- {w}\n" for w in warnings)) if warnings else ""
    excluded_note=("")
    if excluded:
        excluded_note=(f"\n\nEXCLUDED (below top {MAX_MEMO_CANDIDATES}):\n"
                       +", ".join(s.get("name","?") for s in excluded))
    dynamic_max_tokens=min(8192,max(4096,len(summaries)*700))
    prompt=f"""Generate a professional SmartSMBAI Growth Agent shortlist memo.

REGION: {region or "All Regions"}
PREPARED BY: {prepared_by}
DATE: {datetime.now(timezone.utc).strftime("%d %B %Y")}
CANDIDATES: {len(summaries)} (top {len(summaries)} by score)

CANDIDATE DATA:
{candidates_text}{warnings_text}{excluded_note}

Write a structured memo with these sections:
1. EXECUTIVE SUMMARY (2-3 sentences on pool quality for {region})
2. RANKED SHORTLIST (each: name, region, score/35, key strength, key concern, commission model understanding, follow-up questions)
3. RECOMMENDED FIRST HIRE (name + 2-sentence rationale including regional fit)
4. LIVE INTERVIEW GUIDANCE (follow-up questions for ALL {len(summaries)} candidates — group by candidate name)
5. CERTIFICATION STATUS (who has completed the 4-hour cert, who still needs to)
6. PROCESS NOTES (regional pool patterns, missing signals, data warnings if any)

Return ONLY valid JSON:
{{
  "executive_summary": string,
  "ranked_candidates": [{{"rank":int,"name":string,"email":string,"region":string,"score":int,
    "key_strength":string,"key_concern":string,"commission_model_understanding":string,
    "follow_up_questions":[list],"certified":bool}}],
  "recommended_first_hire": string,
  "live_interview_guidance": [list],
  "certification_status": [{{"name":string,"certified":bool,"action_required":string}}],
  "process_notes": string
}}"""
    try:
        resp=client.messages.create(model=MODEL,max_tokens=dynamic_max_tokens,system=SYSTEM_PROMPT,
            messages=[{"role":"user","content":[{"type":"text","text":prompt}]}])
        raw=resp.content[0].text.strip().replace("```json","").replace("```","").strip()
        memo=json.loads(raw)
        memo.update({"region":region or "All Regions","prepared_by":prepared_by,
                     "generated_at":datetime.now(timezone.utc).isoformat(),
                     "candidate_count":len(summaries),
                     "total_shortlisted":len(summaries)+len(excluded),
                     "excluded_candidates":[s.get("name","?") for s in excluded],
                     "single_candidate":False,"data_warnings":warnings})
        log_candidates_in_report([s.get("id","") for s in summaries if s.get("id")],prepared_by,region or "")
        _log("report_generated",actor=prepared_by,object_type="region",
             summary=f"Memo — {region} — {len(summaries)} of {len(summaries)+len(excluded)} candidates"
                     +(f" — {len(warnings)} warning(s)" if warnings else ""))
        return memo
    except Exception as e:
        return {"error":str(e),"region":region,"executive_summary":f"Error: {e}",
                "candidate_count":len(summaries),"total_shortlisted":len(summaries)+len(excluded),
                "excluded_candidates":[s.get("name","?") for s in excluded],"data_warnings":warnings}

def run_shortlist_report(region=None, prepared_by="SmartSMBAI HR", candidate_ids=None) -> dict:
    candidates=get_shortlisted_candidates(region, candidate_ids)
    if not candidates:
        return {"error":"No shortlisted candidates found.","region":region,"candidate_count":0}
    if len(candidates)==1: return _single_candidate_memo(candidates,region,prepared_by)
    return _multi_candidate_memo(candidates,region,prepared_by)

# ── Word document generation ────────────────────────────────────────
def _set_cell_bg(cell, rgb: tuple):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    hex_color="%02X%02X%02X"%rgb
    tcPr=cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")): tcPr.remove(existing)
    shd=OxmlElement("w:shd")
    shd.set(qn("w:fill"),hex_color); shd.set(qn("w:color"),"auto"); shd.set(qn("w:val"),"clear")
    tcPr.append(shd)

def _add_para_border(p, side="bottom", val="single", sz=6, space=1, color="CCCCCC"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
    bdr=OxmlElement(f"w:{side}")
    bdr.set(qn("w:val"),val); bdr.set(qn("w:sz"),str(sz))
    bdr.set(qn("w:space"),str(space)); bdr.set(qn("w:color"),color.lstrip("#"))
    pBdr.append(bdr); pPr.append(pBdr)

def _score_bar(score: int) -> str:
    filled=round(score/35*10); return "█"*filled+"░"*(10-filled)+f"  {score}/35"

def generate_word_doc(memo: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc=Document()
    for section in doc.sections:
        section.top_margin=Cm(2.0); section.bottom_margin=Cm(2.0)
        section.left_margin=Cm(2.5); section.right_margin=Cm(2.5)
    hdr_tbl=doc.add_table(rows=1,cols=2); hdr_tbl.style="Table Grid"; hdr_tbl.autofit=False
    hdr_tbl.columns[0].width=Inches(3.5); hdr_tbl.columns[1].width=Inches(4.0)
    lc=hdr_tbl.cell(0,0); rc=hdr_tbl.cell(0,1)
    _set_cell_bg(lc,NAVY); _set_cell_bg(rc,NAVY)
    lp=lc.paragraphs[0]; lp.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for txt,col in [("Smart",RGBColor(255,255,255)),("SMB",RGBColor(*BLUE)),("AI",RGBColor(*TEAL))]:
        r=lp.add_run(txt); r.font.color.rgb=col; r.font.bold=True; r.font.size=Pt(18)
    rp=rc.paragraphs[0]; rp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    rr=rp.add_run("Productized AI for Small & Medium Businesses")
    rr.font.color.rgb=RGBColor(181,212,244); rr.font.size=Pt(9); rr.font.italic=True
    doc.add_paragraph()
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tr=tp.add_run("GROWTH AGENT SHORTLIST MEMO"); tr.font.bold=True; tr.font.size=Pt(20); tr.font.color.rgb=RGBColor(*NAVY)
    rp2=doc.add_paragraph(); rp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr2=rp2.add_run(f"Region: {memo.get('region','—')}"); rr2.font.size=Pt(14); rr2.font.bold=True; rr2.font.color.rgb=RGBColor(*BLUE)
    mp=doc.add_paragraph(); mp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    mr=mp.add_run(f"Prepared by: {memo.get('prepared_by','')}   ·   Date: {memo.get('generated_at','')[:10]}   ·   Candidates: {memo.get('candidate_count',0)}")
    mr.font.size=Pt(9); mr.font.color.rgb=RGBColor(*GRAY); mr.font.italic=True
    doc.add_paragraph()
    if memo.get("single_candidate"):
        nt=doc.add_table(rows=1,cols=1); nt.style="Table Grid"; nc=nt.cell(0,0)
        _set_cell_bg(nc,(254,243,199)); np_=nc.paragraphs[0]
        nr=np_.add_run("NOTE: Single-candidate memo. No comparative ranking generated.")
        nr.font.size=Pt(9); nr.font.color.rgb=RGBColor(*AMBER)
        doc.add_paragraph()
    def _h(text,sz,color):
        p=doc.add_paragraph(); p.spacing_before=Pt(12); p.spacing_after=Pt(6)
        r=p.add_run(text); r.font.bold=True; r.font.size=Pt(sz); r.font.color.rgb=RGBColor(*color); return p
    _h("1.  Executive Summary",13,NAVY)
    es=doc.add_paragraph(); esr=es.add_run(memo.get("executive_summary","—")); esr.font.size=Pt(10); esr.font.color.rgb=RGBColor(*GRAY)
    doc.add_paragraph()
    _h("2.  Shortlisted Candidates",13,NAVY)
    for cand in memo.get("ranked_candidates",[]):
        name=cand.get("name","—"); score=cand.get("score",0); region=cand.get("region","—")
        ch=doc.add_paragraph()
        cr=ch.add_run(f"#{cand.get('rank','')}  {name}  ({region})")
        cr.font.bold=True; cr.font.size=Pt(12); cr.font.color.rgb=RGBColor(*NAVY)
        sb=doc.add_paragraph()
        sr=sb.add_run(_score_bar(score)); sr.font.size=Pt(10); sr.font.name="Courier New"
        score_color=GREEN if score>=22 else (AMBER if score>=16 else RED)
        sr.font.color.rgb=RGBColor(*score_color)
        cert_p=doc.add_paragraph()
        cert_r=cert_p.add_run(f"Certification: {'✓ Complete' if cand.get('certified') else '⏳ Pending — must complete before client onboarding'}")
        cert_r.font.size=Pt(9); cert_r.font.color.rgb=RGBColor(*(GREEN if cand.get("certified") else AMBER))
        sc_tbl=doc.add_table(rows=3,cols=2); sc_tbl.style="Table Grid"; sc_tbl.autofit=False
        sc_tbl.columns[0].width=Inches(1.4); sc_tbl.columns[1].width=Inches(6.1)
        for ri,(label,field,cell_color) in enumerate([
            ("Key Strength","key_strength",(234,243,222)),
            ("Key Concern","key_concern",(252,235,235)),
            ("Commission Model","commission_model_understanding",(224,234,255))]):
            lbl_cell=sc_tbl.cell(ri,0); val_cell=sc_tbl.cell(ri,1)
            _set_cell_bg(lbl_cell,cell_color); _set_cell_bg(val_cell,cell_color)
            fc=GREEN if ri==0 else (RED if ri==1 else NAVY)
            for cell_obj,text,bold in [(lbl_cell,label,True),(val_cell,cand.get(field,"—"),False)]:
                p=cell_obj.paragraphs[0]; r=p.add_run(text); r.font.size=Pt(9)
                r.font.color.rgb=RGBColor(*fc); r.font.bold=bold
        fu=cand.get("follow_up_questions",[])
        if fu:
            fup=doc.add_paragraph(); fup.add_run("Live interview follow-up questions:").font.bold=True
            for q in fu:
                qp=doc.add_paragraph(style="List Bullet")
                qr=qp.add_run(q); qr.font.size=Pt(9); qr.font.color.rgb=RGBColor(*GRAY)
        doc.add_paragraph()
    if memo.get("recommended_first_hire"):
        _h("3.  Recommended First Hire",13,NAVY)
        rt=doc.add_table(rows=1,cols=1); rt.style="Table Grid"; rc_=rt.cell(0,0)
        _set_cell_bg(rc_,(234,243,222)); rtp=rc_.paragraphs[0]
        rtr=rtp.add_run(memo["recommended_first_hire"]); rtr.font.size=Pt(10); rtr.font.bold=True; rtr.font.color.rgb=RGBColor(*GREEN)
        doc.add_paragraph()
    lig=memo.get("live_interview_guidance",[])
    if lig:
        _h("4.  Live Interview Guidance",13,NAVY)
        for q in lig:
            qp=doc.add_paragraph(style="List Bullet"); qr=qp.add_run(q); qr.font.size=Pt(10); qr.font.color.rgb=RGBColor(*GRAY)
        doc.add_paragraph()
    cert_status=memo.get("certification_status",[])
    if cert_status:
        _h("5.  Certification Status",13,NAVY)
        for cs in cert_status:
            icon="✅" if cs.get("certified") else "⏳"
            p=doc.add_paragraph(); r=p.add_run(f"{icon}  {cs.get('name','—')} — {cs.get('action_required','')}")
            r.font.size=Pt(10); r.font.color.rgb=RGBColor(*(GREEN if cs.get("certified") else AMBER))
        doc.add_paragraph()
    if memo.get("process_notes"):
        _h("6.  Process Notes",13,NAVY)
        p=doc.add_paragraph(); r=p.add_run(memo["process_notes"]); r.font.size=Pt(10); r.font.color.rgb=RGBColor(*GRAY)
    ft=doc.add_paragraph(); ft.paragraph_format.space_before=Pt(12)
    _add_para_border(ft,"top",sz=6,color="CCCCCC")
    ftr=ft.add_run(f"SmartSMBAI  ·  Confidential HR Document  ·  Generated {datetime.now(timezone.utc).strftime('%d %B %Y at %H:%M UTC')}")
    ftr.font.size=Pt(8); ftr.font.color.rgb=RGBColor(*GRAY); ftr.font.italic=True
    buf=BytesIO(); doc.save(buf); return buf.getvalue()

def generate_cv_word_doc(candidate_name: str, region: str, cv_text: str, cover_letter: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc=Document()
    for section in doc.sections:
        section.top_margin=Cm(2.0); section.bottom_margin=Cm(2.0)
        section.left_margin=Cm(2.5); section.right_margin=Cm(2.5)
    hdr_tbl=doc.add_table(rows=1,cols=1); hdr_tbl.style="Table Grid"; hc=hdr_tbl.cell(0,0)
    _set_cell_bg(hc,NAVY); hp=hc.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for txt,col in [("Smart",RGBColor(255,255,255)),("SMB",RGBColor(*BLUE)),("AI",RGBColor(*TEAL))]:
        r=hp.add_run(txt); r.font.color.rgb=col; r.font.bold=True; r.font.size=Pt(16)
    hr3=hp.add_run("   ·   Candidate CV Extract")
    hr3.font.color.rgb=RGBColor(181,212,244); hr3.font.size=Pt(10); hr3.font.italic=True
    doc.add_paragraph()
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tr=tp.add_run(candidate_name); tr.font.bold=True; tr.font.size=Pt(18); tr.font.color.rgb=RGBColor(*NAVY)
    rp=doc.add_paragraph(); rp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=rp.add_run(f"Growth Agent Applicant — {region}"); rr.font.size=Pt(11); rr.font.bold=True; rr.font.color.rgb=RGBColor(*BLUE); rr.font.italic=True
    dp=doc.add_paragraph(); dp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    dr=dp.add_run(f"SmartSMBAI HR  ·  {datetime.now(timezone.utc).strftime('%d %B %Y')}")
    dr.font.size=Pt(9); dr.font.color.rgb=RGBColor(*GRAY); dr.font.italic=True
    doc.add_paragraph()
    def _section_head(text):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
        _add_para_border(p,"bottom",sz=4,color="1A2B5E")
        r=p.add_run(text.upper()); r.font.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(*NAVY)
    if cover_letter and cover_letter.strip():
        _section_head("Cover Letter / Application")
        for para in cover_letter.strip().split("\n"):
            if para.strip():
                p=doc.add_paragraph(para.strip()); p.runs[0].font.size=Pt(10); p.runs[0].font.color.rgb=RGBColor(*GRAY)
        doc.add_paragraph()
    _section_head("CV / Resume")
    cv_content=cv_text.strip() if cv_text and cv_text.strip() else "(No CV text extracted from attachments)"
    for para in cv_content.split("\n"):
        if para.strip():
            p=doc.add_paragraph(para.strip()); p.runs[0].font.size=Pt(10); p.runs[0].font.color.rgb=RGBColor(*GRAY)
    doc.add_paragraph()
    ft=doc.add_paragraph(); _add_para_border(ft,"top",sz=4,color="CCCCCC")
    ftr=ft.add_run("SmartSMBAI HR Suite  ·  Candidate personal data — handle per data retention policy.")
    ftr.font.size=Pt(8); ftr.font.color.rgb=RGBColor(*GRAY); ftr.font.italic=True
    buf=BytesIO(); doc.save(buf); return buf.getvalue()

def get_shortlist_history(region=None, candidate_id=None, limit=200) -> list:
    from database import get_shortlist_history as _get_hist
    return _get_hist(region=region, candidate_id=candidate_id, limit=limit)
